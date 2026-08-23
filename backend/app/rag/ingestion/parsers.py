"""Document parsers.

Each parser turns raw bytes into an ordered list of :class:`ParsedBlock` --
a paragraph-sized unit of text that still knows where it came from (page
number, heading path).  Preserving that provenance through parsing is what
makes citations like "security_policy.pdf - page 4" possible later; a parser
that returns one flat string throws the information away permanently.

SECURITY: file type is decided by *content sniffing*, not by the client-supplied
filename or Content-Type header.  Both are attacker-controlled.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Protocol

from app.core.exceptions import IngestionError
from app.core.logging import get_logger

logger = get_logger("app.rag.parsers")

# Upper bounds that stop a small file from expanding into a very large
# ingestion job ("decompression bomb" style resource exhaustion).
MAX_PAGES = 2000
MAX_EXTRACTED_CHARS = 8_000_000


@dataclass
class ParsedBlock:
    """One contiguous piece of document text plus its provenance."""

    text: str
    page_number: int | None = None
    section: str | None = None
    kind: str = "paragraph"  # paragraph | heading | list | table


@dataclass
class ParsedDocument:
    blocks: list[ParsedBlock] = field(default_factory=list)
    page_count: int = 0
    meta: dict[str, str] = field(default_factory=dict)

    @property
    def char_count(self) -> int:
        return sum(len(b.text) for b in self.blocks)


class Parser(Protocol):
    def parse(self, data: bytes, filename: str) -> ParsedDocument: ...


# ---------------------------------------------------------------------------
# Content sniffing
# ---------------------------------------------------------------------------

_MAGIC = {
    b"%PDF-": "pdf",
    b"PK\x03\x04": "zip",  # DOCX is a ZIP container
}


def sniff_format(data: bytes, declared_extension: str) -> str:
    """Determine the true format of ``data``.

    Returns a canonical extension.  Raises if the bytes clearly contradict the
    declared extension -- a ``.txt`` that is really a ZIP, or a ``.pdf`` that is
    not a PDF, is either a broken client or an attempt to reach a parser that
    was not meant to see this file.
    """
    head = data[:8]
    sniffed: str | None = None
    for magic, kind in _MAGIC.items():
        if head.startswith(magic):
            sniffed = kind
            break

    ext = declared_extension.lower().lstrip(".")
    if ext == "markdown":
        ext = "md"

    if sniffed == "pdf":
        if ext != "pdf":
            raise IngestionError(
                "File content does not match its extension.",
                internal_detail=f"PDF magic bytes with .{ext} extension",
            )
        return "pdf"

    if sniffed == "zip":
        if ext != "docx":
            raise IngestionError(
                "File content does not match its extension.",
                internal_detail=f"ZIP container with .{ext} extension",
            )
        return "docx"

    # No recognised magic: must be a plain-text format.
    if ext in {"txt", "md"}:
        return ext
    raise IngestionError(
        "The file could not be recognised as a supported document.",
        internal_detail=f"no magic match for declared .{ext}",
    )


# ---------------------------------------------------------------------------
# Plain text / Markdown
# ---------------------------------------------------------------------------


def _decode(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    # latin-1 cannot fail, so this is unreachable in practice.
    return data.decode("utf-8", errors="replace")


class TextParser:
    """Plain text: split on blank lines into paragraphs."""

    def parse(self, data: bytes, filename: str) -> ParsedDocument:
        text = _decode(data)[:MAX_EXTRACTED_CHARS]
        blocks = [
            ParsedBlock(text=para.strip())
            for para in re.split(r"\n\s*\n", text)
            if para.strip()
        ]
        return ParsedDocument(blocks=blocks, page_count=1)


_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_MD_FENCE = re.compile(r"^\s*```")


class MarkdownParser:
    """Markdown: track the heading path so chunks inherit their section.

    Fenced code blocks are kept intact -- splitting inside one produces chunks
    that are syntactically meaningless and retrieve badly.
    """

    def parse(self, data: bytes, filename: str) -> ParsedDocument:
        text = _decode(data)[:MAX_EXTRACTED_CHARS]
        blocks: list[ParsedBlock] = []
        heading_stack: list[str] = []
        buffer: list[str] = []
        in_fence = False

        def flush() -> None:
            joined = "\n".join(buffer).strip()
            buffer.clear()
            if joined:
                blocks.append(
                    ParsedBlock(
                        text=joined,
                        section=" > ".join(heading_stack) or None,
                    )
                )

        for line in text.splitlines():
            if _MD_FENCE.match(line):
                in_fence = not in_fence
                buffer.append(line)
                continue
            if in_fence:
                buffer.append(line)
                continue

            heading = _MD_HEADING.match(line)
            if heading:
                flush()
                level = len(heading.group(1))
                title = heading.group(2).strip()
                del heading_stack[level - 1 :]
                heading_stack.append(title)
                blocks.append(
                    ParsedBlock(
                        text=title,
                        section=" > ".join(heading_stack),
                        kind="heading",
                    )
                )
                continue

            if not line.strip():
                flush()
            else:
                buffer.append(line)

        flush()
        return ParsedDocument(blocks=blocks, page_count=1)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


class PdfParser:
    """PDF text extraction with page attribution.

    ``pypdf`` is a pure-Python extractor: it reads the text layer and does not
    OCR.  Scanned PDFs therefore yield little or nothing, which the ingestion
    pipeline surfaces as an explicit failure rather than an empty document.
    """

    def parse(self, data: bytes, filename: str) -> ParsedDocument:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError

        try:
            reader = PdfReader(io.BytesIO(data))
        except PdfReadError as exc:
            raise IngestionError(
                "The PDF could not be read.",
                internal_detail=f"pypdf: {exc}",
            ) from exc

        if reader.is_encrypted:
            # An empty-password decrypt covers PDFs that are "encrypted" only
            # to set permissions, which is common for policy documents.
            try:
                if reader.decrypt("") == 0:
                    raise IngestionError(
                        "Password-protected PDFs are not supported.",
                        internal_detail="pypdf decrypt returned 0",
                    )
            except IngestionError:
                raise
            except Exception as exc:
                raise IngestionError(
                    "Password-protected PDFs are not supported.",
                    internal_detail=f"pypdf decrypt failed: {exc}",
                ) from exc

        pages = reader.pages
        if len(pages) > MAX_PAGES:
            raise IngestionError(
                f"The PDF exceeds the {MAX_PAGES}-page limit.",
                internal_detail=f"page_count={len(pages)}",
            )

        blocks: list[ParsedBlock] = []
        total_chars = 0
        for index, page in enumerate(pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception as exc:  # a single corrupt page must not abort
                logger.warning(
                    "pdf_page_extraction_failed",
                    extra={"page": index, "error": type(exc).__name__},
                )
                continue
            for para in re.split(r"\n\s*\n", text):
                para = para.strip()
                if not para:
                    continue
                blocks.append(ParsedBlock(text=para, page_number=index))
                total_chars += len(para)
            if total_chars > MAX_EXTRACTED_CHARS:
                logger.warning("pdf_truncated", extra={"page": index})
                break

        return ParsedDocument(
            blocks=blocks,
            page_count=len(pages),
            meta={"producer": str(reader.metadata.producer or "")}
            if reader.metadata
            else {},
        )


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


class DocxParser:
    """DOCX paragraphs and tables, with heading styles mapped to sections."""

    def parse(self, data: bytes, filename: str) -> ParsedDocument:
        import docx
        from docx.opc.exceptions import PackageNotFoundError

        try:
            document = docx.Document(io.BytesIO(data))
        except (PackageNotFoundError, KeyError, ValueError) as exc:
            raise IngestionError(
                "The DOCX file could not be read.",
                internal_detail=f"python-docx: {exc}",
            ) from exc

        blocks: list[ParsedBlock] = []
        heading_stack: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name or "") if paragraph.style else ""
            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                del heading_stack[level - 1 :]
                heading_stack.append(text)
                blocks.append(
                    ParsedBlock(
                        text=text, section=" > ".join(heading_stack), kind="heading"
                    )
                )
            else:
                blocks.append(
                    ParsedBlock(text=text, section=" > ".join(heading_stack) or None)
                )

        # Tables carry a lot of policy content (leave allowances, limits...).
        # Rendering rows as pipe-delimited lines keeps them searchable.
        for table in document.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                blocks.append(
                    ParsedBlock(
                        text="\n".join(rows),
                        section=" > ".join(heading_stack) or None,
                        kind="table",
                    )
                )

        return ParsedDocument(blocks=blocks, page_count=1)


_PARSERS: dict[str, Parser] = {
    "pdf": PdfParser(),
    "docx": DocxParser(),
    "md": MarkdownParser(),
    "txt": TextParser(),
}


def parse_document(data: bytes, filename: str, extension: str) -> ParsedDocument:
    """Parse ``data`` using the parser for its *sniffed* format."""
    fmt = sniff_format(data, extension)
    parser = _PARSERS.get(fmt)
    if parser is None:  # pragma: no cover - sniff_format already restricts this
        raise IngestionError(
            "Unsupported document type.", internal_detail=f"format={fmt}"
        )
    parsed = parser.parse(data, filename)
    if not parsed.blocks:
        raise IngestionError(
            "No readable text was found in the document. "
            "Scanned images are not supported without OCR.",
            internal_detail=f"format={fmt} produced zero blocks",
        )
    return parsed
